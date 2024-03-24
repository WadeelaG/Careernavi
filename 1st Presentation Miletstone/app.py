from flask import Flask, request, render_template, redirect, url_for, session, send_file
from cover_letter import CoverLetterGenerate  # Adjusted import
from docx import Document
from resume import ResumeGenerate
from fpdf import FPDF
import os
import io

app = Flask(__name__)
app.secret_key = 'your_secret_key'  # Set to a unique secret value for session management


@app.route('/index.html')
def index():
    return render_template('index.html')

@app.route('/service.html')
def service():
    return render_template('service.html')

@app.route('/about.html')
def about():
    return render_template('about.html')

@app.route('/login-form.html')
def login():
    return render_template('login-form.html')

@app.route('/Register-form.html')
def Register():
    return render_template('Register-form.html')

@app.route('/coverletter-form.html')
def coverletter_form():
    return render_template('coverletter-form.html')

@app.route('/resume-form.html')
def resume_form():
    return render_template('resume-form.html')

@app.route('/blog-single.html')
def blog_single():
    return render_template('blog-single.html')

@app.route('/contact.html')
def contact():
    return render_template('contact.html')

@app.route('/PerInfo.html')
def PerInfo():
    return render_template('PerInfo.html')

@app.route('/Dummy 1 Profile.html')
def profile():
    return render_template('Dummy 1 Profile.html')

@app.route('/project.html')
def project():
    return render_template('project.html')


@app.route('/cover_letter', methods=['GET', 'POST'])
def cover_letter():
    if request.method == 'POST':
        # Extract form data
        applicant_name = f"{request.form['firstName']} {request.form['lastName']}"
        job_title = request.form['jobTitle']
        company_name = request.form['companyName']
        phone_number = request.form['phoneNumber']
        email_address = request.form['emailAddress']
        experience = request.form['experience']
        skills= request.form['skills'] 
        
        # Create an instance of the cover letter generator
        generator = CoverLetterGenerate(applicant_name, job_title, company_name, phone_number, email_address, experience,skills)
        cover_letter = generator.generate_cover_letter()

        # Store the cover letter in the session
        session['cover_letter'] = cover_letter

        # Redirect to the cover_letter_output route
        return redirect(url_for('cover_letter_output'))

    else:
        # If the method is GET, show the form to the user
        return render_template('coverletter-form.html')

@app.route('/cover_letter_output', methods=['GET'])
def cover_letter_output():
    # Retrieve the generated cover letter from the session
    cover_letter = session.get('cover_letter', None)
    return render_template('coverletter-output.html', cover_letter=cover_letter)


def save_cover_letter_as_docx(cover_letter, filename="Cover_Letter.docx"):
    doc = Document()
    doc.add_paragraph(cover_letter)
    doc.save(filename)

def save_cover_letter_as_pdf(cover_letter, filename="Cover_Letter.pdf"):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, cover_letter)
    pdf.output(filename)

@app.route('/download_cover_letter/docx',methods=['GET'])
def download_cover_letter_docx():
    cover_letter = session.get('cover_letter', 'No cover letter generated.')
    filename = "Cover_Letter.docx"
    save_cover_letter_as_docx(cover_letter, filename)
    return send_file(filename, as_attachment=True)

@app.route('/download_cover_letter/pdf',methods=['GET'])
def download_cover_letter_pdf():
    cover_letter = session.get('cover_letter', 'No cover letter generated.')
    filename = "Cover_Letter.pdf"
    save_cover_letter_as_pdf(cover_letter, filename)
    return send_file(filename, as_attachment=True)



@app.route('/resume-content', methods=['GET', 'POST'])
def resume():
    if request.method == 'POST':
        # Extract form data
        fullname = request.form['fullname']
        phoneNumber = request.form['phoneNumber']
        emailAddress = request.form['emailAddress']
        address = request.form['address']
        city = request.form['city']
        state = request.form['state']
        zipCode = request.form['zipCode']
        objective = request.form['objective']
        education = request.form['education']
        experience = request.form['experience']
        technicalSkills = request.form['technicalSkills']
        softSkills = request.form['softSkills']
        certifications = request.form['certifications']
        projects = request.form['projects']
        languages = request.form['languages']
        additionalInformation = request.form['additionalInformation']
        linkdinprofile = request.form['linkdinprofile']
        professionalSummary = request.form['professionalSummary']

        # Create an instance of the resume generator
        generator = ResumeGenerate(fullname, phoneNumber, emailAddress, address, city, state, zipCode, objective, education, experience, technicalSkills, softSkills, certifications, projects, languages, additionalInformation, linkdinprofile, professionalSummary)
        resume_content = generator.generate_resume()

        # Store the resume in the session
        session['resume_content'] = resume_content

        # Redirect to the resume_output route
        return redirect(url_for('resume_output'))

    else:
        # If the method is GET, show the form to the user
        return render_template('resume-form.html')

@app.route('/resume_output', methods=['GET'])
def resume_output():
    # Retrieve the generated resume from the session
    resume_content = session.get('resume_content', None)
    return render_template('resume-output.html', resume_content=resume_content)

# Function to save the resume as DOCX
def save_resume_as_docx(resume_content, filename="Resume.docx"):
    doc = Document()
    doc.add_paragraph(resume_content)
    doc.save(filename)

# Function to save the resume as PDF
def save_resume_as_pdf(resume_content, filename="Resume.pdf"):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, resume_content)
    pdf.output(filename)

@app.route('/download_resume/docx', methods=['GET'])
def download_resume_docx():
    resume_content = session.get('resume_content', 'No resume generated.')
    filename = "Resume.docx"
    save_resume_as_docx(resume_content, filename)
    return send_file(filename, as_attachment=True)

@app.route('/download_resume/pdf', methods=['GET'])
def download_resume_pdf():
    resume_content = session.get('resume_content', 'No resume generated.')
    filename = "Resume.pdf"
    save_resume_as_pdf(resume_content, filename)
    return send_file(filename, as_attachment=True)




if __name__ == '__main__':
    app.run(debug=True)
